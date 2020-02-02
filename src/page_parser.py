import argparse
import json
from docx import Document
import cv2
import numpy as np
from line_parser import LineParser


class PageParser(LineParser):
    def __init__(self, hashes, CHARS_PER_LINE):
        LineParser.__init__(self, hashes)
        self.CHARS_PER_LINE = CHARS_PER_LINE
    
    # Returns a list of lines that belong to the same page
    def parse_page(self, document, show = False):
        # denotes text that doesn't fit in a line an must be printed in the next line
        leftover = ''
        # list that stores image of each line
        lines = []

        # In every iteration, generate a line and wrap the leftover text to next line
        for para in document.paragraphs:
            text = leftover + str(para.text)
            image, leftover = self.parse_line_constrained(text, self.CHARS_PER_LINE)
            lines.append(image)

        # print all leftover text
        while leftover != '':
            text = leftover
            image, leftover = self.parse_line_constrained(text, self.CHARS_PER_LINE)
            lines.append(image)

        if len(lines) == 0:
            print('Empty document!')
            return lines

        page = lines[0]
        for i in range(1, len(lines)):
            page = np.vstack((page, lines[i]))
        if show:
            this.show(page)
        return lines

    # Parse page - add page breaks when max no of lines in page is exceeded
    # fill incomplete page with lines of whitespaces
    # Uses parse_page(), returns a list of images of each page
    def parse_pages_constrained(self, document, LINES_PER_PAGE, show = False):
        lines = self.parse_page(document, show = False)
        totalLines = len(lines)
        totalPages = (totalLines // LINES_PER_PAGE)

        # Work-around to https://stackoverflow.com/questions/19951816/python-changes-to-my-copy-variable-affect-the-original-variable
        # Instead of appending to a list of final images with temporary variable
        # Initialize list to required size and modify value at each index 
        finalImages = [[[i]] for i in range(0, totalPages + 1)]
        
        remainderLines = totalLines % LINES_PER_PAGE
        blankLines = LINES_PER_PAGE - (remainderLines)

        # Compute blank lines
        blanks = []
        blankline = self.parse_line(' ' * self.CHARS_PER_LINE)
        for i in range(0, blankLines):
            blanks.append(blankline)
        
        # Print all complete pages
        pageIndex = 0
        for i in range(0, totalPages):
            starting = True
            for j in range(i*LINES_PER_PAGE, (i+1)*LINES_PER_PAGE):
                if starting:
                    finalImages[pageIndex] = lines[j]
                    starting = False
                else:
                    finalImages[pageIndex] = np.vstack((finalImages[pageIndex], lines[j]))
            pageIndex += 1
        
        # Print incomplete page
        starting = True
        for i in range(0, remainderLines):
            if starting:
                finalImages[pageIndex] = lines[totalPages + i]
                starting = False
            else:
                finalImages[pageIndex] = np.vstack((finalImages[pageIndex], lines[totalPages + i]))
            

        # Print blanks at end of incomplete page
        for blank in blanks:
            finalImages[pageIndex] = np.vstack((finalImages[pageIndex], blank))

        # show all pages
        if show:
            for finalImage in finalImages:
                self.show('window', finalImage)
        
        return finalImages

    # Generates page from a list of lines
    def generate_page(line_list):
        if len(line_list) > 0:
            page = line_list[0]
            for i in range(1, len(line_list)):
                page = np.vstack((page, line_list[i]))
            return page
        else:
            print('Empty image list!')
            return [[]]



                
def main():
    line_parser = line_parser.LineParser(hashes)
    leftover = ''   
    lines = []      
    for para in document.paragraphs:
        text = leftover + para.text
        image, leftover = line_parser.parse_line_constrained(text, CHARS_PER_LINE)
        lines.append(image)
    while leftover != '':
        text = leftover
        image, leftover = line_parser.parse_line_constrained(text, CHARS_PER_LINE)
        lines.append(image)
    page = lines[0]
    for i in range(1, len(lines)):
        page = np.vstack((page, lines[i]))
    line_parser.show('window', page)
                
            

if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Output page for docx page')
    parser.add_argument('document_path', type=str, nargs=1)
    args = parser.parse_args()
    CHARS_PER_LINE = 54
    document = Document('test.docx')
    with open('hashes.json') as f:
        hashes = json.load(f)
    main()